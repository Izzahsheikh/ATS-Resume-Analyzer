import os
import re
import json
import string
import hashlib
from concurrent.futures import ThreadPoolExecutor, as_completed

import io

import streamlit as st
from dotenv import load_dotenv
import fitz  # PyMuPDF
import docx  # python-docx
import numpy as np
from langchain_google_genai import ChatGoogleGenerativeAI
from sentence_transformers import SentenceTransformer
from langchain_core.prompts import PromptTemplate

load_dotenv()

# --------------------------------------------------------------------------
# PAGE CONFIG
# --------------------------------------------------------------------------
st.set_page_config(
    page_title="Izra | AI Resume Analyzer",
    page_icon="assets/izra_logo_cropped.png",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --------------------------------------------------------------------------
# CUSTOM UI THEME
# --------------------------------------------------------------------------
st.markdown("""
<style>

/* =========================================================
   GLOBAL
   ========================================================= */
.stApp {
    background-color: #F7F6F4;
}

.main .block-container {
    padding-top: 2rem;
    padding-bottom: 3rem;
}

/* =========================================================
   SCROLLBAR — FLAT BLACK, NO HOVER, NO GREY (applies everywhere)
   ========================================================= */
* {
    scrollbar-width: auto !important;
    scrollbar-color: #000000 transparent !important;
}

*::-webkit-scrollbar {
    width: 14px !important;
    height: 14px !important;
}

*::-webkit-scrollbar-track,
*::-webkit-scrollbar-track:hover,
*::-webkit-scrollbar-track:active {
    background: transparent !important;
}

*::-webkit-scrollbar-thumb,
*::-webkit-scrollbar-thumb:hover,
*::-webkit-scrollbar-thumb:active {
    background: #000000 !important;
    border: none !important;
    border-radius: 10px !important;
    min-height: 60px !important;
}

*::-webkit-scrollbar-corner {
    background: transparent !important;
}

*::-webkit-scrollbar-button {
    display: none !important;
    width: 0 !important;
    height: 0 !important;
}

/* =========================================================
   SIDEBAR
   ========================================================= */
section[data-testid="stSidebar"] {
    background-color: #EFEEEC !important;
    border-right: 1px solid #DDD9D7;
}

section[data-testid="stSidebar"] h1,
section[data-testid="stSidebar"] h2,
section[data-testid="stSidebar"] h3 {
    color: #541F49 !important;
}

section[data-testid="stSidebar"] p {
    color: #292529 !important;
}

section[data-testid="stSidebar"] .stMarkdown p {
    color: #6F6A6D !important;
    font-size: 15px !important;
    line-height: 1.5 !important;
}

/* =========================================================
   MAIN HEADINGS
   ========================================================= */
h1 { color: #541F49 !important; font-size: 32px !important; font-weight: 700 !important; }
h2 { color: #541F49 !important; font-size: 22px !important; font-weight: 650 !important; }
h3 { color: #292529 !important; font-size: 18px !important; font-weight: 600 !important; }

/* =========================================================
   THRESHOLD VALUE BOX
   ========================================================= */
section[data-testid="stSidebar"] div[data-testid="stTextInput"] > div {
    background-color: #FFFFFF !important;
    border-radius: 7px !important;
}

section[data-testid="stSidebar"] div[data-testid="stTextInput"] input {
    background-color: #FFFFFF !important;
    color: #292529 !important;
    -webkit-text-fill-color: #292529 !important;
    text-align: center !important;
    font-size: 18px !important;
    font-weight: 600 !important;
    border: 1px solid #D8D4D1 !important;
    border-radius: 7px !important;
    color-scheme: light !important;
}

section[data-testid="stSidebar"] div[data-testid="stTextInput"] input:focus {
    background-color: #FFFFFF !important;
    color: #292529 !important;
    -webkit-text-fill-color: #292529 !important;
    border-color: #541F49 !important;
    box-shadow: 0 0 0 1px #541F49 !important;
}

/* =========================================================
   THRESHOLD MINUS / PLUS BUTTONS
   ========================================================= */
section[data-testid="stSidebar"] .stButton > button {
    min-height: 40px !important;
    padding: 0 !important;
    background-color: #541F49 !important;
    border: 1px solid #541F49 !important;
    border-radius: 7px !important;
    color: #FFFFFF !important;
    font-size: 22px !important;
    font-weight: 600 !important;
}

section[data-testid="stSidebar"] .stButton > button p,
section[data-testid="stSidebar"] .stButton > button span {
    color: #FFFFFF !important;
}

section[data-testid="stSidebar"] .stButton > button:hover {
    background-color: #6B3A5E !important;
    border-color: #6B3A5E !important;
    color: #FFFFFF !important;
}

section[data-testid="stSidebar"] .stButton > button:focus,
section[data-testid="stSidebar"] .stButton > button:focus-visible,
section[data-testid="stSidebar"] .stButton > button:active {
    background-color: #541F49 !important;
    border-color: #541F49 !important;
    color: #FFFFFF !important;
    box-shadow: none !important;
}

/* =========================================================
   MAIN BUTTONS
   ========================================================= */
.stButton > button {
    background-color: #6B3A5E !important;
    border: 1px solid #6B3A5E !important;
    border-radius: 8px !important;
    min-height: 48px !important;
    padding: 0 24px !important;
    color: #FFFFFF !important;
    font-size: 19px !important;
    font-weight: 600 !important;
    transition: background-color 0.15s ease, border-color 0.15s ease !important;
}

.stButton > button p, .stButton > button span { color: #FFFFFF !important; }
.stButton > button:hover { background-color: #541F49 !important; border-color: #541F49 !important; }
.stButton > button:focus, .stButton > button:focus-visible, .stButton > button:active {
    background-color: #6B3A5E !important;
    border-color: #6B3A5E !important;
    color: #FFFFFF !important;
    box-shadow: none !important;
}

/* =========================================================
   FILE UPLOADER — equal height for both columns
   ========================================================= */
[data-testid="stFileUploader"] {
    background-color: #F0EFEC !important;
    border: 1px solid #D8D4D1 !important;
    border-radius: 10px !important;
    padding: 6px !important;
    margin-top: 4px !important;
}

[data-testid="stFileUploaderDropzone"] {
    background-color: #F0EFEC !important;
    border: 2px dashed #C0B4BE !important;
    border-radius: 8px !important;
    padding: 20px 16px !important;
}

/* Dropzone instruction text */
[data-testid="stFileUploaderDropzoneInstructions"] {
    font-size: 15px !important;
    color: #6F6A6D !important;
}

[data-testid="stFileUploaderDropzoneInstructions"] span {
    font-size: 15px !important;
    color: #6F6A6D !important;
}

/* Drag active state */
[data-testid="stFileUploaderDropzone"]:hover,
[data-testid="stFileUploaderDropzone"][data-dragging="true"] {
    border-color: #541F49 !important;
    background-color: #EDE0EC !important;
}

/* =========================================================
   FILE UPLOADER LABEL
   ========================================================= */
[data-testid="stFileUploader"] label {
    color: #292529 !important;
    -webkit-text-fill-color: #292529 !important;
    font-size: 15px !important;
    font-weight: 600 !important;
    line-height: 1.5 !important;
}

[data-testid="stFileUploader"] [data-testid="stTooltipHoverTarget"] {
    opacity: 0 !important;
    transition: opacity 0.15s ease !important;
}

[data-testid="stFileUploader"]:hover [data-testid="stTooltipHoverTarget"] {
    opacity: 1 !important;
}

[data-testid="stFileUploader"] small {
    color: #292529 !important;
    -webkit-text-fill-color: #292529 !important;
    font-size: 14px !important;
    line-height: 1.5 !important;
}

/* =========================================================
   UPLOAD BUTTON (Browse / Upload)
   ========================================================= */
[data-testid="stFileUploader"] button {
    background-color: #6B3A5E !important;
    border: 1px solid #6B3A5E !important;
    border-radius: 7px !important;
    min-height: 42px !important;
    padding: 0 18px !important;
    color: #FFFFFF !important;
    font-size: 16px !important;
    font-weight: 600 !important;
}

[data-testid="stFileUploader"] button span,
[data-testid="stFileUploader"] button p { color: #FFFFFF !important; }
[data-testid="stFileUploader"] button:hover { background-color: #541F49 !important; border-color: #541F49 !important; }
[data-testid="stFileUploader"] button:focus,
[data-testid="stFileUploader"] button:focus-visible,
[data-testid="stFileUploader"] button:active {
    background-color: #6B3A5E !important;
    border-color: #6B3A5E !important;
    color: #FFFFFF !important;
    box-shadow: none !important;
}

[data-testid="stFileUploaderDeleteBtn"] button {
    background-color: transparent !important;
    border: none !important;
    min-height: unset !important;
    padding: 2px !important;
    color: #541F49 !important;
    font-size: 13px !important;
}

/* =========================================================
   ADD MORE FILES (+) BUTTON
   ========================================================= */
[data-testid="stFileUploaderFileList"] + div button,
[data-testid="stFileUploaderFileList"] ~ div button {
    background-color: #EDE0EC !important;
    border: 1.5px dashed #C9A8C4 !important;
    border-radius: 8px !important;
    min-height: 42px !important;
    min-width: 58px !important;
    padding: 0 14px !important;
    color: #541F49 !important;
    font-size: 20px !important;
    font-weight: 700 !important;
    margin-top: 0 !important;
    box-shadow: none !important;
    cursor: pointer !important;
}

[data-testid="stFileUploaderFileList"] + div button:hover,
[data-testid="stFileUploaderFileList"] ~ div button:hover {
    background-color: #DCC7D9 !important;
    border-color: #541F49 !important;
    color: #541F49 !important;
}

[data-testid="stFileUploaderFileList"] + div button:focus,
[data-testid="stFileUploaderFileList"] + div button:focus-visible,
[data-testid="stFileUploaderFileList"] ~ div button:focus,
[data-testid="stFileUploaderFileList"] ~ div button:focus-visible {
    background-color: #EDE0EC !important;
    border-color: #541F49 !important;
    color: #541F49 !important;
    box-shadow: 0 0 0 2px rgba(84, 31, 73, 0.15) !important;
}

[data-testid="stFileUploaderFileList"] + div button svg,
[data-testid="stFileUploaderFileList"] + div button span,
[data-testid="stFileUploaderFileList"] + div button p,
[data-testid="stFileUploaderFileList"] ~ div button svg,
[data-testid="stFileUploaderFileList"] ~ div button span,
[data-testid="stFileUploaderFileList"] ~ div button p {
    color: #541F49 !important;
    fill: #541F49 !important;
    -webkit-text-fill-color: #541F49 !important;
}

/* =========================================================
   DROPZONE
   ========================================================= */
[data-testid="stFileUploaderDropzone"] {
    min-height: 80px !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
}

[data-testid="stFileUploaderDropzoneInstructions"] {
    visibility: visible !important;
    opacity: 1 !important;
}

/* =========================================================
   JD UPLOADER — hide "Add files" (+) button entirely.
   ========================================================= */
.st-key-jd_upload_container [data-testid="stFileUploader"] button[aria-label="Add files"] {
    display: none !important;
}

.st-key-jd_upload_container [data-testid="stFileUploaderFileList"] + div {
    margin-top: 0 !important;
}

/* =========================================================
   CV UPLOADER — scrollable file list, always-visible thumb
   ========================================================= */
.st-key-cv_upload_container [data-testid="stFileUploaderFileList"] {
    max-height: 200px !important;
    overflow-y: scroll !important;
    overflow-x: hidden !important;
}

.st-key-cv_upload_container [data-testid="stFileUploaderFileList"]::-webkit-scrollbar {
    width: 14px !important;
}

.st-key-cv_upload_container [data-testid="stFileUploaderFileList"]::-webkit-scrollbar-track {
    background: transparent !important;
}

.st-key-cv_upload_container [data-testid="stFileUploaderFileList"]::-webkit-scrollbar-thumb {
    background: #000000 !important;
    border-radius: 10px !important;
    min-height: 60px !important;
}

/* =========================================================
   METRICS
   ========================================================= */
[data-testid="stMetricValue"] {
    color: #541F49 !important;
    font-size: 28px !important;
    font-weight: 700 !important;
}

[data-testid="stMetricLabel"] {
    color: #6F6A6D !important;
    font-size: 16px !important;
}

/* =========================================================
   EXPANDERS / CARDS
   ========================================================= */
[data-testid="stExpander"] {
    background-color: #FFFFFF !important;
    border: 1px solid #DDD9D7 !important;
    border-radius: 8px !important;
}

[data-testid="stExpander"] summary {
    font-size: 16px !important;
    font-weight: 600 !important;
    background-color: #FFFFFF !important;
    color: #541F49 !important;
    -webkit-text-fill-color: #541F49 !important;
    border-radius: 8px !important;
}

[data-testid="stExpander"] summary p,
[data-testid="stExpander"] summary span {
    color: #541F49 !important;
    -webkit-text-fill-color: #541F49 !important;
}

[data-testid="stExpander"] details[open] summary {
    background-color: #541F49 !important;
    color: #FFFFFF !important;
    -webkit-text-fill-color: #FFFFFF !important;
}

[data-testid="stExpander"] details[open] summary p,
[data-testid="stExpander"] details[open] summary span {
    color: #FFFFFF !important;
    -webkit-text-fill-color: #FFFFFF !important;
}

[data-testid="stExpander"] summary:focus { outline: none !important; }
[data-testid="stExpander"] summary:hover {
    background-color: #F3F1F2 !important;
    color: #541F49 !important;
    -webkit-text-fill-color: #541F49 !important;
}

/* =========================================================
   DIVIDERS
   ========================================================= */
hr { border-color: #DDD9D7 !important; }

/* =========================================================
   CAPTIONS
   ========================================================= */
[data-testid="stCaptionContainer"] {
    color: #6F6A6D !important;
    font-size: 15px !important;
}

/* =========================================================
   ALERTS
   ========================================================= */
[data-testid="stAlert"] {
    border-radius: 8px !important;
    font-size: 16px !important;
}

/* =========================================================
   GENERAL MARKDOWN
   ========================================================= */
.stMarkdown { font-size: 16px; }

/* =========================================================
   RESULT CARD TEXT
   ========================================================= */
[data-testid="stVerticalBlockBorderWrapper"] { background-color: #FFFFFF !important; }

[data-testid="stVerticalBlockBorderWrapper"] p,
[data-testid="stVerticalBlockBorderWrapper"] span,
[data-testid="stVerticalBlockBorderWrapper"] h3 {
    color: #292529 !important;
    -webkit-text-fill-color: #292529 !important;
}

[data-testid="stVerticalBlockBorderWrapper"] .stMarkdown div {
    color: #292529 !important;
    -webkit-text-fill-color: #292529 !important;
}

[data-testid="stExpander"] p,
[data-testid="stExpander"] span,
[data-testid="stExpander"] li { color: #292529 !important; }

[data-testid="stExpander"] strong { color: #292529 !important; }

/* =========================================================
   PROGRESS TEXT
   ========================================================= */
[data-testid="stProgress"] p { color: #292529 !important; }
[data-testid="stProgress"] span { color: #292529 !important; }

[data-testid="stVerticalBlockBorderWrapper"] [data-testid="stMarkdownContainer"] h3 {
    color: #292529 !important;
    -webkit-text-fill-color: #292529 !important;
    opacity: 1 !important;
}

/* =========================================================
   STEP LABELS
   ========================================================= */
.step-label {
    display: inline-block;
    font-size: 13px;
    font-weight: 800;
    letter-spacing: 0.14em;
    text-transform: uppercase;
    color: #541F49;
    background-color: #EDE0EC;
    border: 1.5px solid #C9A8C4;
    border-radius: 5px;
    padding: 5px 14px;
    margin-bottom: 12px;
}

/* =========================================================
   NO SHORTLIST MESSAGE
   ========================================================= */
.custom-no-shortlist {
    background-color: #FFF1EB !important;
    border: 1px solid #F26B4F !important;
    border-radius: 8px !important;
    padding: 12px 16px !important;
    display: flex !important;
    align-items: center !important;
    gap: 8px !important;
}

.custom-no-shortlist-icon {
    color: #F26B4F !important;
    -webkit-text-fill-color: #F26B4F !important;
    font-size: 18px !important;
    font-weight: 700 !important;
}

.custom-no-shortlist-text {
    color: #B94A32 !important;
    -webkit-text-fill-color: #B94A32 !important;
    font-size: 15px !important;
    font-weight: 600 !important;
    opacity: 1 !important;
}

/* =========================================================
   RECOMMENDATION SECTION
   ========================================================= */
.rec-section {
    background: #FFFFFF;
    border: 1.5px solid #D4B8D0;
    border-radius: 12px;
    padding: 24px 28px;
    margin-bottom: 24px;
}

.rec-eyebrow {
    font-size: 11px;
    font-weight: 700;
    letter-spacing: 0.12em;
    text-transform: uppercase;
    color: #541F49;
    margin-bottom: 4px;
}

.rec-heading {
    font-size: 22px;
    font-weight: 700;
    color: #292529;
    margin: 0 0 6px 0;
}

.rec-subtext {
    font-size: 14px;
    color: #6F6A6D;
    margin: 0 0 18px 0;
}

.rec-candidate-list {
    display: flex;
    flex-wrap: wrap;
    gap: 10px;
    margin-top: 4px;
}

.rec-candidate-pill {
    display: inline-flex;
    align-items: center;
    gap: 7px;
    background: #F0E8EF;
    border: 1px solid #D4B8D0;
    border-radius: 6px;
    padding: 6px 14px;
    font-size: 14px;
    font-weight: 600;
    color: #541F49;
}

.rec-none {
    background: #FFF1EB;
    border: 1px solid #F26B4F;
    border-radius: 8px;
    padding: 12px 16px;
    font-size: 14px;
    font-weight: 600;
    color: #B94A32;
}

/* =========================================================
   CV UPLOADER — "Upload More" button label
   ========================================================= */
[data-testid="stFileUploader"] button[aria-label="Add files"] {
    width: 105px !important;
    height: 42px !important;
    font-size: 0 !important;
}

[data-testid="stFileUploader"] button[aria-label="Add files"] svg {
    display: none !important;
}

[data-testid="stFileUploader"] button[aria-label="Add files"]::before {
    content: "Upload More" !important;
    font-size: 15px !important;
    font-weight: 600 !important;
    line-height: 1 !important;
    white-space: nowrap !important;
}

/* =========================================================
   THRESHOLD INPUT — main area (not sidebar)
   ========================================================= */
div[data-testid="stTextInput"] > div {
    background-color: #FFFFFF !important;
    border-radius: 7px !important;
}

div[data-testid="stTextInput"] input {
    background-color: #FFFFFF !important;
    color: #292529 !important;
    -webkit-text-fill-color: #292529 !important;
    text-align: center !important;
    font-size: 18px !important;
    font-weight: 600 !important;
    border: 1px solid #D8D4D1 !important;
    border-radius: 7px !important;
    color-scheme: light !important;
}

div[data-testid="stTextInput"] input:focus {
    background-color: #FFFFFF !important;
    color: #292529 !important;
    -webkit-text-fill-color: #292529 !important;
    border-color: #541F49 !important;
    box-shadow: 0 0 0 1px #541F49 !important;
}

</style>
""", unsafe_allow_html=True)


# --------------------------------------------------------------------------
# HEADER
# --------------------------------------------------------------------------
logo_col, title_col = st.columns([0.7, 5])

with logo_col:
    st.image("assets/izra_logo_cropped.png", width=150)

with title_col:
    st.markdown(
        "<h1 style='margin-bottom:0;'>AI Resume Analyzer</h1>",
        unsafe_allow_html=True
    )
    st.markdown(
        "<p style='color:#6F6A6D; font-size:16px; margin-top:4px;'>"
        "Find the right candidates faster with AI-powered screening."
        "</p>",
        unsafe_allow_html=True
    )

# --------------------------------------------------------------------------
# HELPERS
# --------------------------------------------------------------------------
def normalize_extracted_text(text: str) -> str:
    text = re.sub(r'(?<=[a-z])(?=[A-Z])', ' ', text)
    text = re.sub(r'(?<=[A-Za-z])(?=\d)', ' ', text)
    text = re.sub(r'(?<=\d)(?=[A-Za-z])', ' ', text)
    text = re.sub(r'(?<=[,/])(?=\S)', ' ', text)
    text = re.sub(r'[ \t]+', ' ', text)
    return text


def extract_text_from_uploaded(uploaded_file) -> str:
    suffix = os.path.splitext(uploaded_file.name)[1].lower()
    if suffix == ".pdf":
        file_bytes = uploaded_file.read()
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text() + "\n"
        doc.close()
        return normalize_extracted_text(text)
    elif suffix == ".docx":
        file_bytes = uploaded_file.read()
        document = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        text = "\n".join(paragraphs)
        return normalize_extracted_text(text)
    else:
        raw = uploaded_file.read().decode("utf-8", errors="ignore")
        return normalize_extracted_text(raw)


def parse_json_response(text):
    if isinstance(text, list):
        text = "".join(
            block.get("text", "") if isinstance(block, dict) else str(block)
            for block in text
        )
    cleaned = re.sub(r"```json|```", "", text).strip()
    match = re.search(r"\{.*\}", cleaned, re.DOTALL)
    if not match:
        raise ValueError("No JSON object found in model output")
    return json.loads(match.group(0))


def get_llm():
    return ChatGoogleGenerativeAI(model="gemini-3.5-flash", temperature=0)


# --------------------------------------------------------------------------
# STEP 1: Extract required skills from JD
# --------------------------------------------------------------------------
JD_SKILLS_PROMPT = PromptTemplate(
    input_variables=["job_description"],
    template="""
You are a technical recruiter. Read the job description below and extract the specific
technical skills, tools, frameworks, and technologies required for this role.

JOB DESCRIPTION:
{job_description}

Respond with ONLY a valid JSON object (no markdown, no extra text):
{{
  "required_skills": ["skill1", "skill2", "skill3"]
}}

Be specific — include exact tool/framework names like "LangChain", "RAG", "Docker", "PyTorch".
Do NOT include soft skills like "communication" or "teamwork".
""",
)


@st.cache_data(show_spinner=False)
def extract_jd_skills(jd_text: str) -> list:
    prompt = JD_SKILLS_PROMPT.format(job_description=jd_text)
    response = get_llm().invoke(prompt)
    data = parse_json_response(response.content)
    return [s.lower() for s in data.get("required_skills", [])]


@st.cache_resource
def get_embedding_model():
    return SentenceTransformer("all-MiniLM-L6-v2")


# --------------------------------------------------------------------------
# STEP 2: Semantic Embeddings
# --------------------------------------------------------------------------
@st.cache_data(show_spinner=False)
def generate_embeddings(texts: tuple) -> np.ndarray:
    model = get_embedding_model()
    embeddings = model.encode(
        list(texts),
        batch_size=32,
        convert_to_numpy=True,
        normalize_embeddings=True,
        show_progress_bar=False
    )
    return embeddings


def compute_semantic_similarity(jd_embedding, resume_embeddings):
    similarities = np.dot(resume_embeddings, jd_embedding)
    return [round(float(score) * 100) for score in similarities]


# --------------------------------------------------------------------------
# STEP 3: Deterministic scoring
# --------------------------------------------------------------------------
def compute_skill_match(resume_text, required_skills):
    if not required_skills:
        return 0, [], []

    chunks = [line.strip() for line in resume_text.splitlines() if len(line.strip()) >= 15]
    if not chunks:
        return 0, [], []

    model = get_embedding_model()
    skill_embeddings = model.encode(required_skills, batch_size=32, convert_to_numpy=True, normalize_embeddings=True, show_progress_bar=False)
    chunk_embeddings = model.encode(chunks, batch_size=32, convert_to_numpy=True, normalize_embeddings=True, show_progress_bar=False)

    similarities = np.dot(skill_embeddings, chunk_embeddings.T)
    matched, missing = [], []
    MATCH_THRESHOLD = 0.50

    for skill, skill_sims in zip(required_skills, similarities):
        if float(np.max(skill_sims)) >= MATCH_THRESHOLD:
            matched.append(skill)
        else:
            missing.append(skill)

    score = round(len(matched) / len(required_skills) * 100)
    return score, matched, missing


def score_education(resume_text):
    text = resume_text.lower()
    if any(w in text for w in ["phd", "doctorate"]): return 100
    elif any(w in text for w in ["master", "msc", "ms ", "m.s", "mba"]): return 85
    elif any(w in text for w in ["bachelor", "bsc", "b.s", "b.e", "beng", "degree"]): return 70
    elif any(w in text for w in ["diploma", "associate"]): return 50
    return 30


def score_experience(resume_text, jd_text):
    resume_years = re.findall(r"(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience|exp)", resume_text.lower())
    jd_years = re.findall(r"(\d+)\+?\s*years?\s*(?:of\s*)?(?:experience|exp)", jd_text.lower())
    if resume_years and jd_years:
        candidate_exp = max(int(y) for y in resume_years)
        required_exp = max(int(y) for y in jd_years)
        if candidate_exp >= required_exp: return 100
        return round((candidate_exp / required_exp) * 100)
    text = resume_text.lower()
    if any(w in text for w in ["senior", "lead", "principal", "head of", "manager"]): return 85
    elif any(w in text for w in ["engineer", "developer", "analyst", "apprentice"]): return 65
    elif any(w in text for w in ["junior", "intern", "trainee", "fresher", "graduate"]): return 40
    return 50


def score_projects(resume_text):
    text = resume_text.lower()
    score = 0
    if "project" in text: score += 50
    if any(w in text for w in ["github", "gitlab", "bitbucket"]): score += 25
    if any(w in text for w in ["deployed", "production", "live", "published"]): score += 25
    return score


def score_formatting(resume_text):
    lines = [l.strip() for l in resume_text.splitlines() if l.strip()]
    score = 50
    if len(lines) > 20: score += 15
    if any(w in resume_text.lower() for w in ["experience", "education", "skills", "projects"]): score += 20
    if any(w in resume_text.lower() for w in ["@", "email", "phone", "linkedin"]): score += 15
    return min(score, 100)


def score_resume_deterministic(resume_text, jd_text, required_skills, semantic_score, resume_embedding, skill_embeddings):
    skill_score, matched, missing = compute_skill_match(resume_text, required_skills)
    combined_skill = round(skill_score * 0.7 + semantic_score * 0.3)
    experience_score = score_experience(resume_text, jd_text)
    education_score = score_education(resume_text)
    projects_score = score_projects(resume_text)
    formatting_score = score_formatting(resume_text)
    ats_score = round(
        combined_skill   * 0.40 +
        experience_score * 0.25 +
        education_score  * 0.15 +
        projects_score   * 0.12 +
        formatting_score * 0.08
    )
    return {
        "ats_score": ats_score,
        "final_score": ats_score,
        "skill_match": combined_skill,
        "experience_score": experience_score,
        "education_score": education_score,
        "projects_score": projects_score,
        "formatting_score": formatting_score,
        "matched_skills": [s.title() for s in matched],
        "missing_skills": [s.title() for s in missing],
    }


# --------------------------------------------------------------------------
# SESSION STATE
# --------------------------------------------------------------------------
if "shortlist_threshold" not in st.session_state:
    st.session_state.shortlist_threshold = 70
if "threshold_text" not in st.session_state:
    st.session_state.threshold_text = "70"


def decrease_threshold():
    new_value = max(0, st.session_state.shortlist_threshold - 1)
    st.session_state.shortlist_threshold = new_value
    st.session_state.threshold_text = str(new_value)


def increase_threshold():
    new_value = min(100, st.session_state.shortlist_threshold + 1)
    st.session_state.shortlist_threshold = new_value
    st.session_state.threshold_text = str(new_value)


def update_threshold():
    try:
        value = int(st.session_state.threshold_text)
        if 0 <= value <= 100:
            st.session_state.shortlist_threshold = value
        else:
            st.session_state.threshold_text = str(st.session_state.shortlist_threshold)
    except ValueError:
        st.session_state.threshold_text = str(st.session_state.shortlist_threshold)


# --------------------------------------------------------------------------
# SIDEBAR — branding + screening settings + threshold
# --------------------------------------------------------------------------
with st.sidebar:

    # ── Branding ──────────────────────────────────────────────────────────
    st.image("assets/izra_logo_cropped.png", width=110)
    st.markdown(
        "<p style='font-size:13px; color:#6F6A6D; margin-top:2px; margin-bottom:16px;'>"
        "AI-powered candidate screening</p>",
        unsafe_allow_html=True
    )
    st.markdown("---")

    # ── Screening Settings ─────────────────────────────────────────────────
    st.markdown(
        "<h2 style='margin:20px 0 6px 0;'>Screening Settings</h2>",
        unsafe_allow_html=True
    )
    st.markdown(
        "<p style='font-size:15px; color:#6F6A6D; margin:0 0 28px 0; line-height:1.5;'>"
        "Set your score cutoff for shortlisting.</p>",
        unsafe_allow_html=True
    )

    # ── Step 3 / 3 — Shortlist Threshold ─────────────────────────────────
    st.markdown(
        "<div style='margin-bottom:10px;'><span class='step-label'>STEP 3 / 3</span></div>",
        unsafe_allow_html=True
    )
    st.markdown(
        "<p style='font-size:20px; font-weight:700; color:#292529 !important; -webkit-text-fill-color:#292529 !important; margin:0 0 8px 0; line-height:1.4;'>Shortlist Threshold</p>",
        unsafe_allow_html=True
    )

    minus_col, value_col, plus_col = st.columns([1, 1.5, 1])

    with minus_col:
        st.button("−", key="threshold_minus", use_container_width=True, on_click=decrease_threshold)

    with value_col:
        st.text_input("Threshold", key="threshold_text", label_visibility="collapsed", on_change=update_threshold)

    with plus_col:
        st.button("＋", key="threshold_plus", use_container_width=True, on_click=increase_threshold)

    shortlist_threshold = st.session_state.shortlist_threshold

    st.markdown(
        f"<p style='color:#6F6A6D; font-size:13px; margin-top:8px;'>"
        f"Candidates scoring <b>{shortlist_threshold}%</b> or higher will be shortlisted.</p>",
        unsafe_allow_html=True
    )
    st.markdown("---")


# --------------------------------------------------------------------------
# SECTION HEADINGS ROW
# --------------------------------------------------------------------------
heading_left, heading_right = st.columns(2, gap="large")

with heading_left:
    st.markdown("<h2 style='margin-top:20px; margin-bottom:6px;'>Start Screening</h2>", unsafe_allow_html=True)
    st.markdown(
        "<p style='color:#6F6A6D; font-size:16px; margin-bottom:28px; line-height:1.6;'>"
        "Upload your JD and candidate CVs, set threshold then click <b>Analyze Candidates</b>.</p>",
        unsafe_allow_html=True
    )

with heading_right:
    st.markdown("<p style='color:#6F6A6D; font-size:16px; margin-bottom:28px; margin-top:20px; line-height:1.6;'>&nbsp;</p>", unsafe_allow_html=True)


# --------------------------------------------------------------------------
# UPLOAD COLUMNS
# --------------------------------------------------------------------------
jd_col, cv_col = st.columns(2, gap="large")

# --------------------------------------------------------------------------
# JOB DESCRIPTION  (STEP 1 / 3)
# --------------------------------------------------------------------------
with jd_col:
    st.markdown(
        """
        <div style="margin-bottom: 10px;">
            <span class="step-label">STEP 1 / 3</span>
        </div>
        """,
        unsafe_allow_html=True
    )
    st.markdown(
        "<p style='font-size:20px; font-weight:700; color:#292529; margin:0 0 8px 0; line-height:1.4;'>📋 Job Description</p>",
        unsafe_allow_html=True
    )
    st.markdown(
        "<p style='font-size:15px; color:#6F6A6D; margin:0 0 14px 0; line-height:1.6;'>"
        "Upload a single JD file — PDF, DOCX, or TXT. Drag &amp; drop or click <b>Upload</b>.</p>",
        unsafe_allow_html=True
    )

    with st.container(key="jd_upload_container"):
        jd_file = st.file_uploader(
            "Upload Job Description",
            type=["pdf", "txt", "docx"],
            key="jd_upload",
            help="Upload a PDF, TXT or DOCX job description."
        )


# --------------------------------------------------------------------------
# CANDIDATE CVs  (STEP 2 / 3)
# --------------------------------------------------------------------------
with cv_col:
    st.markdown(
        """
        <div style="margin-bottom: 10px;">
            <span class="step-label">STEP 2 / 3</span>
        </div>
        """,
        unsafe_allow_html=True
    )
    st.markdown(
        "<p style='font-size:20px; font-weight:700; color:#292529; margin:0 0 8px 0; line-height:1.4;'>👥 Candidate CVs</p>",
        unsafe_allow_html=True
    )
    st.markdown(
        "<p style='font-size:15px; color:#6F6A6D; margin:0 0 14px 0; line-height:1.6;'>"
        "Upload one or more candidate PDF resumes.</p>",
        unsafe_allow_html=True
    )

    with st.container(key="cv_upload_container"):
        resume_files = st.file_uploader(
            "Upload Candidate CVs",
            type=["pdf"],
            accept_multiple_files=True,
            key="resume_upload",
            help="Select multiple candidate resumes. After uploading, click + to add more files."
        )

    if resume_files:
        st.markdown(
            f"<p style='color:#292529; font-size:14px; margin-top:6px;'>"
            f"{len(resume_files)} resume{'s' if len(resume_files) != 1 else ''} added</p>",
            unsafe_allow_html=True
        )


# --------------------------------------------------------------------------
# ANALYZE BUTTON
# --------------------------------------------------------------------------
st.markdown("<br>", unsafe_allow_html=True)
button_col1, button_col2, button_col3 = st.columns([1.5, 2, 1.5])

with button_col2:
    analyze_clicked = st.button(
        "✦ Analyze Candidates",
        type="primary",
        use_container_width=True
    )

# Ensure shortlist_threshold is always defined before the pipeline
shortlist_threshold = st.session_state.shortlist_threshold

# --------------------------------------------------------------------------
# MAIN PIPELINE
# --------------------------------------------------------------------------
if analyze_clicked:
    if not jd_file:
        st.error("⚠️ Step 1 is incomplete — please upload a Job Description before analyzing candidates.")
        st.stop()
    if not resume_files:
        st.error("⚠️ Step 2 is incomplete — please add at least one Candidate CV.")
        st.stop()

    # ── JD processing ─────────────────────────────────────────────────────
    with st.spinner("Reading job description..."):
        jd_text = extract_text_from_uploaded(jd_file)

    if not jd_text.strip():
        st.error("Couldn't extract text from the job description.")
        st.stop()

    with st.spinner("Extracting required skills from job description..."):
        required_skills = extract_jd_skills(jd_text)

    # ── PHASE 1: Parallel PDF extraction ──────────────────────────────────
    progress = st.progress(0.0, text="Extracting resume text in parallel...")
    resume_texts: dict = {}
    failed: list = []

    def extract_one(rf):
        text = extract_text_from_uploaded(rf)
        return rf.name, text

    with ThreadPoolExecutor(max_workers=min(len(resume_files), 8)) as executor:
        futures = {executor.submit(extract_one, rf): rf.name for rf in resume_files}
        done = 0
        for future in as_completed(futures):
            name = futures[future]
            try:
                fname, text = future.result()
                if text.strip():
                    resume_texts[fname] = text
                else:
                    st.warning(f"Skipping {fname} — no extractable text.")
                    failed.append(fname)
            except Exception as e:
                st.warning(f"Failed to read {name}: {e}")
                failed.append(name)
            done += 1
            progress.progress(done / len(resume_files) * 0.3, text=f"Extracted {done}/{len(resume_files)} resumes...")

    if not resume_texts:
        st.error("No resumes could be read.")
        st.stop()

    # ── PHASE 2: Local Semantic Embeddings ────────────────────────────────
    progress.progress(0.35, text="Generating local semantic embeddings...")
    names_ordered = list(resume_texts.keys())
    texts_ordered = [resume_texts[name] for name in names_ordered]

    jd_embedding = generate_embeddings((jd_text,))[0]
    resume_embeddings = generate_embeddings(tuple(texts_ordered))
    skill_embeddings = generate_embeddings(tuple(required_skills))

    semantic_scores = compute_semantic_similarity(jd_embedding, resume_embeddings)
    semantic_map = dict(zip(names_ordered, semantic_scores))
    resume_embedding_map = dict(zip(names_ordered, resume_embeddings))

    # ── PHASE 3: Parallel deterministic scoring ────────────────────────────
    progress.progress(0.4, text="Scoring resumes in parallel...")
    results: list = []

    def score_one(name):
        text = resume_texts[name]
        result = score_resume_deterministic(
            resume_text=text,
            jd_text=jd_text,
            required_skills=required_skills,
            semantic_score=semantic_map[name],
            resume_embedding=resume_embedding_map[name],
            skill_embeddings=skill_embeddings,
        )
        result["candidate_name"] = name
        return result

    with ThreadPoolExecutor(max_workers=min(len(resume_texts), 8)) as executor:
        futures = {executor.submit(score_one, n): n for n in names_ordered}
        done = 0
        for future in as_completed(futures):
            try:
                results.append(future.result())
            except Exception as e:
                name = futures[future]
                st.warning(f"Scoring failed for {name}: {e}")
            done += 1
            progress.progress(0.4 + done / len(resume_texts) * 0.3, text=f"Scored {done}/{len(resume_texts)} resumes...")

    if not results:
        st.error("No resumes could be scored.")
        st.stop()

    progress.empty()

    # ── PHASE 4: Rank ──────────────────────────────────────────────────────
    results.sort(key=lambda r: r["final_score"], reverse=True)

    shortlisted_names = [
        r["candidate_name"]
        for r in results
        if r["final_score"] >= shortlist_threshold
    ]

    # ── RECOMMENDATION SECTION ─────────────────────────────────────────────
    st.markdown("<div id='results-anchor'></div>", unsafe_allow_html=True)
    st.divider()

    st.markdown("<span class='step-label' style='margin-bottom:8px; display:inline-block;'>RESULTS</span>", unsafe_allow_html=True)

    if shortlisted_names:
        pills_html = "".join(
            f"<span class='rec-candidate-pill'>✓ {name}</span>"
            for name in shortlisted_names
        )
        st.markdown(
            f"""
            <div class="rec-section">
                <div class="rec-eyebrow">Recommendation</div>
                <div class="rec-heading">Recommended to Shortlist</div>
                <div class="rec-subtext">
                    These candidates meet or exceed the {shortlist_threshold}% screening threshold
                    and are recommended for further review.
                </div>
                <div class="rec-candidate-list">
                    {pills_html}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )
    else:
        st.markdown(
            f"""
            <div class="rec-section">
                <div class="rec-eyebrow">Recommendation</div>
                <div class="rec-heading">No Candidates Recommended</div>
                <div class="rec-subtext">
                    No candidates met the {shortlist_threshold}% shortlist threshold.
                    Try lowering the threshold in the sidebar settings.
                </div>
                <div class="rec-none">
                    ⚠ No candidates to shortlist at the current threshold.
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

    # ── CANDIDATE RANKING ──────────────────────────────────────────────────
    st.markdown(
        """
        <div class="rec-section" style="margin-top:8px;">
            <div class="rec-eyebrow">Rankings</div>
            <div class="rec-heading">🏆 Candidate Ranking</div>
            <div class="rec-subtext">All candidates ranked by ATS score from highest to lowest.</div>
        </div>
        """,
        unsafe_allow_html=True
    )

    for rank, r in enumerate(results, start=1):
        shortlisted = r["final_score"] >= shortlist_threshold

        badge_html = (
            "<span style='background:#E8F5E9; color:#2E7D32; border:1px solid #A5D6A7; "
            "border-radius:6px; padding:5px 14px; font-size:14px; font-weight:700;'>✓ Shortlisted</span>"
            if shortlisted else
            "<span style='background:#FDECEC; color:#C62828; border:1px solid #EF9A9A; "
            "border-radius:6px; padding:5px 14px; font-size:14px; font-weight:700;'>✕ Not Shortlisted</span>"
        )

        rank_medal = {1: "🥇", 2: "🥈", 3: "🥉"}.get(rank, f"#{rank}")

        score_color = "#2E7D32" if r['final_score'] >= shortlist_threshold else "#C62828"

        st.markdown(
            f"""
            <div style="
                background:#FFFFFF;
                border: 1.5px solid #D4B8D0;
                border-radius:12px;
                padding:20px 24px;
                margin-bottom:14px;
                display:flex;
                align-items:center;
                justify-content:space-between;
                flex-wrap:wrap;
                gap:12px;
            ">
                <div style="display:flex; align-items:center; gap:16px; flex:1; min-width:200px;">
                    <div style="
                        font-size:22px; font-weight:800; color:#541F49;
                        background:#EDE0EC; border:1.5px solid #C9A8C4;
                        border-radius:8px; padding:6px 14px; white-space:nowrap;
                    ">{rank_medal}</div>
                    <div>
                        <div style="font-size:16px; font-weight:700; color:#292529;">{r['candidate_name']}</div>
                        <div style="font-size:13px; color:#6F6A6D; margin-top:2px;">ATS Score</div>
                    </div>
                </div>
                <div style="display:flex; align-items:center; gap:20px; flex-wrap:wrap;">
                    <div style="
                        font-size:28px; font-weight:800; color:{score_color};
                        min-width:70px; text-align:center;
                    ">{r['final_score']}%</div>
                    {badge_html}
                </div>
            </div>
            """,
            unsafe_allow_html=True
        )

        with st.expander(f"View Details — {r['candidate_name']}"):
            c1, c2, c3, c4, c5 = st.columns(5)
            c1.metric("Skill Match", f"{r.get('skill_match', 0)}%")
            c2.metric("Experience", f"{r.get('experience_score', 0)}%")
            c3.metric("Education", f"{r.get('education_score', 0)}%")
            c4.metric("Projects", f"{r.get('projects_score', 0)}%")
            c5.metric("Formatting", f"{r.get('formatting_score', 0)}%")

            d1, d2 = st.columns(2)

            with d1:
                st.markdown("**✅ Matched Skills**")
                for skill in r.get("matched_skills", []):
                    st.markdown(f"- {skill}")

            with d2:
                st.markdown("**❌ Missing Skills**")
                for skill in r.get("missing_skills", []):
                    st.markdown(f"- {skill}")

    # ── Smooth-scroll down to the results section ───────────────────────────
    st.markdown(
        """
        <script>
        setTimeout(function() {
            const el = window.parent.document.getElementById('results-anchor');
            if (el) {
                el.scrollIntoView({ behavior: 'smooth', block: 'start' });
            }
        }, 300);
        </script>
        """,
        unsafe_allow_html=True
    )